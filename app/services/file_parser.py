"""
Handles extracting raw text from uploaded resume files.
Supports PDF and DOCX. JD is usually pasted as plain text,
but we support file upload for it too (same parser).
"""
import io
import pdfplumber
from docx import Document
from fastapi import UploadFile, HTTPException


async def extract_text_from_upload(file: UploadFile) -> str:
    """Reads an UploadFile and returns extracted plain text."""
    filename = (file.filename or "").lower()
    content = await file.read()

    if filename.endswith(".pdf"):
        return _extract_pdf_text(content)
    elif filename.endswith(".docx"):
        return _extract_docx_text(content)
    elif filename.endswith(".txt"):
        return content.decode("utf-8", errors="ignore")
    else:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {filename}. Use PDF, DOCX, or TXT.",
        )


def _extract_pdf_text(content: bytes) -> str:
    text_chunks = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    text = "\n".join(text_chunks).strip()
    if not text:
        raise HTTPException(
            status_code=422,
            detail="Could not extract text from PDF (it may be scanned/image-based).",
        )
    return text


def _extract_docx_text(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    text = "\n".join(paragraphs).strip()
    if not text:
        raise HTTPException(status_code=422, detail="Could not extract text from DOCX.")
    return text